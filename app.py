import streamlit as st
import google.generativeai as genai
import PyPDF2
import docx
import requests
import stripe
import json
import re
from datetime import datetime, timedelta
from fpdf import FPDF
from io import BytesIO
from docx import Document
from PIL import Image
import numpy as np

# ---------------------------
# 1. Configuration & Secrets
# ---------------------------
st.set_page_config(page_title="AI Career Intelligence", page_icon="📈", layout="wide")

# Required secrets
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
STRIPE_SECRET_KEY = st.secrets["STRIPE_SECRET_KEY"]
STRIPE_PRICE_ID_PREMIUM_MONTHLY = st.secrets["STRIPE_PRICE_ID_PREMIUM_MONTHLY"]
STRIPE_PRICE_ID_PREMIUM_LIFETIME = st.secrets["STRIPE_PRICE_ID_PREMIUM_LIFETIME"]
STRIPE_PRICE_ID_PRO_MONTHLY = st.secrets["STRIPE_PRICE_ID_PRO_MONTHLY"]
STRIPE_PRICE_ID_PRO_LIFETIME = st.secrets["STRIPE_PRICE_ID_PRO_LIFETIME"]
APP_URL = st.secrets["APP_URL"]
PREMIUM_UNLOCK_CODE = st.secrets["PREMIUM_UNLOCK_CODE"].strip()
PRO_UNLOCK_CODE = st.secrets["PRO_UNLOCK_CODE"].strip()

# Adzuna credentials (clean version - no warning)
ADZUNA_APP_ID = st.secrets["ADZUNA_APP_ID"]
ADZUNA_APP_KEY = st.secrets["ADZUNA_API_KEY"]  # Your exact secret name

# Optional: RapidAPI key for JSearch fallback
RAPIDAPI_KEY = st.secrets.get("RAPIDAPI_KEY", "")

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("gemini-2.0-flash")
stripe.api_key = STRIPE_SECRET_KEY

COUNTRY_MAP = {
    "United States": "us",
    "United Kingdom": "gb",
    "Canada": "ca",
    "Australia": "au",
    "Germany": "de",
    "France": "fr",
    "India": "in",
    "South Africa": "za",
    "Nigeria": "ng",
    "Kenya": "ke",
    "Botswana": "bw",
    "Ghana": "gh",
    "Other": "other"
}

# ---------------------------
# 2. Custom CSS
# ---------------------------
st.markdown("""
<style>
body { background-color: #f8fafc; font-family: 'Inter', sans-serif; color: #0f172a; }
.block-container { padding-top: 2rem; padding-bottom: 2rem; }
.hero {
    text-align: center;
    padding: 2rem 1rem;
    background: linear-gradient(135deg, #f0f9ff, #e6f0ff);
    border-radius: 30px;
    margin-bottom: 2rem;
}
.hero h1 { font-size: 2.5rem; font-weight: 700; color: #0f172a; margin-bottom: 0.5rem; }
.hero p { font-size: 1.2rem; color: #475569; margin-bottom: 1rem; }
.hero-badge { display: inline-block; background-color: #6C63FF; color: white; padding: 0.3rem 1rem; border-radius: 40px; font-size: 0.8rem; margin-bottom: 1rem; }
.metric-card {
    background: white;
    border-radius: 20px;
    padding: 1.2rem;
    box-shadow: 0 4px 12px rgba(0,0,0,0.05);
    text-align: center;
    border: 1px solid #e2e8f0;
}
.metric-value { font-size: 2rem; font-weight: 700; color: #1e293b; }
.metric-label { font-size: 0.85rem; color: #64748b; margin-top: 0.25rem; }
.action-card {
    background: white;
    border-radius: 24px;
    padding: 1.5rem;
    box-shadow: 0 8px 20px rgba(0,0,0,0.05);
    height: 100%;
    border: 1px solid #e2e8f0;
}
.action-title { font-size: 1.3rem; font-weight: 700; margin-bottom: 1rem; color: #0f172a; }
.job-container {
    background: white;
    border-radius: 20px;
    padding: 1.2rem;
    margin-bottom: 1rem;
    border: 1px solid #e2e8f0;
}
.stButton > button {
    border-radius: 40px;
    background: linear-gradient(90deg, #4A90E2, #6C63FF);
    color: white;
    font-weight: 600;
    border: none;
    padding: 0.5rem 1rem;
}
.stButton > button:hover {
    transform: translateY(-1px);
    box-shadow: 0 4px 12px rgba(108,99,255,0.3);
}
.footer { text-align: center; margin-top: 3rem; padding-top: 1.5rem; border-top: 1px solid #e2e8f0; color: #64748b; font-size: 0.8rem; }
@media (prefers-color-scheme: dark) {
    body { background-color: #0f172a; color: #e2e8f0; }
    .hero { background: linear-gradient(135deg, #1e293b, #0f172a); }
    .hero h1 { color: #f1f5f9; }
    .hero p { color: #cbd5e1; }
    .metric-card, .action-card, .job-container, .pricing-card { background-color: #1e293b; border-color: #334155; }
    .metric-value { color: #f1f5f9; }
    .metric-label { color: #94a3b8; }
    .action-title { color: #f1f5f9; }
    .footer { border-top-color: #334155; color: #94a3b8; }
    .stButton > button { background: linear-gradient(90deg, #3b82f6, #8b5cf6); }
}
</style>
""", unsafe_allow_html=True)

# ---------------------------
# 3. Session State
# ---------------------------
if "premium" not in st.session_state:
    st.session_state.premium = False
if "pro" not in st.session_state:
    st.session_state.pro = False
if "page" not in st.session_state:
    st.session_state.page = "intro"
if "cv_text" not in st.session_state:
    st.session_state.cv_text = ""
if "analysis" not in st.session_state:
    st.session_state.analysis = None
if "target_roles" not in st.session_state:
    st.session_state.target_roles = []
if "primary_role" not in st.session_state:
    st.session_state.primary_role = ""
if "jobs" not in st.session_state:
    st.session_state.jobs = []
if "match_scores" not in st.session_state:
    st.session_state.match_scores = {}
if "saved_jobs" not in st.session_state:
    st.session_state.saved_jobs = []
if "generated_cv" not in st.session_state:
    st.session_state.generated_cv = ""
if "cover_letter_for_job" not in st.session_state:
    st.session_state.cover_letter_for_job = None

# Handle Stripe success redirects
if "success_premium_monthly" in st.query_params:
    st.session_state.premium = True
    st.query_params.clear()
    st.session_state.page = "workspace"
if "success_premium_lifetime" in st.query_params:
    st.session_state.premium = True
    st.query_params.clear()
    st.session_state.page = "workspace"
if "success_pro_monthly" in st.query_params:
    st.session_state.pro = True
    st.query_params.clear()
    st.session_state.page = "workspace"
if "success_pro_lifetime" in st.query_params:
    st.session_state.pro = True
    st.query_params.clear()
    st.session_state.page = "workspace"

# ---------------------------
# 4. Helper Functions
# ---------------------------
def extract_text_from_file(uploaded_file):
    if uploaded_file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(uploaded_file)
        return "".join(page.extract_text() for page in reader.pages)
    elif uploaded_file.name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        return "\n".join(para.text for para in doc.paragraphs)
    else:
        return uploaded_file.read().decode("utf-8")

def clean_json_response(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    start_obj = text.find("{")
    end_obj = text.rfind("}")
    start_arr = text.find("[")
    end_arr = text.rfind("]")
    obj_candidate = text[start_obj:end_obj + 1] if start_obj != -1 and end_obj != -1 else ""
    arr_candidate = text[start_arr:end_arr + 1] if start_arr != -1 and end_arr != -1 else ""
    if obj_candidate and arr_candidate:
        return obj_candidate if len(obj_candidate) >= len(arr_candidate) else arr_candidate
    if obj_candidate:
        return obj_candidate
    if arr_candidate:
        return arr_candidate
    return text

def analyze_cv(cv_text, full=False):
    prompt = f"""
    Analyze this CV as a professional recruiter. Return ONLY valid JSON.

    Required fields:
    - strength_score: 0-100
    - ats_score: 0-100
    - interview_likelihood: "Low","Moderate","High"
    - recruiter_verdict: one sentence explaining the scores
    - experience_level: "Entry","Mid","Senior"
    - target_roles: list of 2-3 specific job titles
    - top_strengths: list of 2-3
    - top_weaknesses: list of 2-3

    If full == true, also include:
    - missing_keywords: list of 4-8
    - rewrite_suggestions: list of 3-5

    CV:
    {cv_text[:10000]}
    """
    response = model.generate_content(prompt + f"\nFull: {full}")
    raw = clean_json_response(response.text)
    try:
        return json.loads(raw)
    except:
        return {
            "strength_score": 50,
            "ats_score": 50,
            "interview_likelihood": "Moderate",
            "recruiter_verdict": "Unable to analyze – please try again.",
            "experience_level": "Mid",
            "target_roles": ["Business Analyst"],
            "top_strengths": ["Error parsing response"],
            "top_weaknesses": ["Error parsing response"]
        }

@st.cache_data(show_spinner=False, ttl=3600)
def analyze_cv_cached(cv_text, full=False):
    return analyze_cv(cv_text, full)

def get_interview_percentage(likelihood):
    mapping = {"Low": "0-20%", "Moderate": "30-60%", "High": "65-85%"}
    return mapping.get(likelihood, "30-60%")

@st.cache_data(ttl=3600)
def analyze_cover_letter_full(letter_text, target_role):
    prompt = f"""
    Evaluate this cover letter for a {target_role} position.
    Return ONLY valid JSON:
    {{
        "alignment_score": 0-100,
        "personalization_score": 0-100,
        "impact_score": 0-100,
        "structure_score": 0-100,
        "overall_score": 0-100,
        "verdict": "one sentence",
        "missing_elements": ["element1", "element2", "element3"]
    }}
    Cover letter:
    {letter_text[:4000]}
    """
    response = model.generate_content(prompt)
    raw = clean_json_response(response.text)
    try:
        return json.loads(raw)
    except:
        return {"overall_score": 50, "verdict": "Unable to evaluate", "missing_elements": []}

@st.cache_data(ttl=3600)
def review_cover_letter_basic(letter_text, target_role):
    prompt = f"""
    Review this cover letter for a {target_role} position.
    Return JSON: {{"overall_score": 0-100, "verdict": "one sentence", "top_missing": "single element"}}
    Cover letter: {letter_text[:4000]}
    """
    response = model.generate_content(prompt)
    raw = clean_json_response(response.text)
    try:
        return json.loads(raw)
    except:
        return {"overall_score": 50, "verdict": "Unable to review", "top_missing": "Improve alignment"}

@st.cache_data(ttl=3600)
def generate_cover_letter(cv_text, target_role, company_name=""):
    company_text = f" for {company_name}" if company_name else ""
    prompt = f"""
    Write a professional cover letter for a {target_role} position{company_text}.
    Base it on the candidate's CV below. DO NOT invent experience.
    Return ONLY the cover letter as plain text (250-350 words).
    CV:
    {cv_text[:8000]}
    """
    response = model.generate_content(prompt)
    return response.text

@st.cache_data(ttl=3600)
def generate_job_query(cv_text):
    prompt = f"""
    Extract the single best job title from this CV.
    Return ONLY one title.
    No explanation.
    CV: {cv_text[:6000]}
    """
    response = model.generate_content(prompt)
    return response.text.strip()

def deduplicate_jobs(jobs):
    seen = set()
    unique = []
    for job in jobs:
        key = f"{job.get('title', '')}_{job.get('company', '')}".lower()
        if key not in seen:
            seen.add(key)
            unique.append(job)
    return unique

def parse_adzuna_date(date_str):
    """Robust date parser that handles multiple Adzuna formats."""
    if not date_str:
        return None
    date_str = str(date_str).strip()

    formats = [
        "%Y-%m-%d",
        "%Y-%m-%dT%H:%M:%SZ",
        "%Y-%m-%dT%H:%M:%S.%fZ",
        "%Y-%m-%d %H:%M:%S",
    ]

    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except Exception:
            pass

    # Fallback: try first 10 chars as YYYY-MM-DD
    try:
        return datetime.strptime(date_str[:10], "%Y-%m-%d")
    except Exception:
        return None

def get_jobs_from_adzuna(query, country_code, location_refine, limit=5):
    """Enhanced Adzuna fetcher with robust date parsing and relaxed filtering."""
    url = f"https://api.adzuna.com/v1/api/jobs/{country_code}/search/1"
    params = {
        "app_id": ADZUNA_APP_ID,
        "app_key": ADZUNA_APP_KEY,
        "results_per_page": limit * 2,
        "what": query
    }
    if location_refine and location_refine.strip():
        params["where"] = location_refine.strip()
    
    log_info = {"status": None, "error": None, "raw_count": 0, "active_count": 0}
    
    try:
        resp = requests.get(url, params=params, timeout=10)
        log_info["status"] = resp.status_code
        if resp.status_code != 200:
            log_info["error"] = f"HTTP {resp.status_code}: {resp.text[:200]}"
            return [], log_info
        
        data = resp.json()
        jobs = data.get("results", [])
        log_info["raw_count"] = len(jobs)
        
        formatted = []
        for j in jobs:
            company = j.get("company", {})
            company_name = company.get("display_name", "Unknown") if isinstance(company, dict) else str(company) if company else "Unknown"
            created = j.get("created")
            closing_date = j.get("closing_date")
            date_display = "📅 Date not specified"
            if closing_date:
                date_display = f"📅 Closing: {closing_date}"
            elif created:
                date_display = f"📅 Posted: {created}"
            formatted.append({
                "title": j.get("title", "Untitled"),
                "company": company_name,
                "location": location_refine or country_code.upper(),
                "url": j.get("redirect_url", "#"),
                "description": j.get("description", ""),
                "date_display": date_display,
                "closing_date": closing_date,
                "created": created,
                "is_expired": False
            })
        
        today = datetime.now().date()
        cutoff = datetime.now() - timedelta(days=90)  # Relaxed to 90 days
        
        active = []
        for job in formatted:
            keep = False
            close_date = parse_adzuna_date(job.get("closing_date"))
            created_date = parse_adzuna_date(job.get("created"))
            
            if close_date:
                if close_date.date() >= today:
                    keep = True
            elif created_date:
                if created_date >= cutoff:
                    keep = True
            else:
                # No parsable date → keep the job rather than discarding
                keep = True
            
            if keep:
                active.append(job)
        
        log_info["active_count"] = len(active)
        unique_jobs = deduplicate_jobs(active)[:limit]
        return unique_jobs, log_info
        
    except Exception as e:
        log_info["error"] = str(e)
        return [], log_info

def get_jobs_from_jsearch(query, country_code, location_refine, limit=5):
    """Fallback job search using JSearch API (RapidAPI)."""
    if not RAPIDAPI_KEY:
        return [], {"error": "RAPIDAPI_KEY not configured"}
    
    url = "https://jsearch.p.rapidapi.com/search"
    country_param = country_code.upper()
    querystring = {
        "query": f"{query} {location_refine}" if location_refine else query,
        "page": "1",
        "num_pages": "1",
        "country": country_param,
        "date_posted": "month"
    }
    headers = {
        "X-RapidAPI-Key": RAPIDAPI_KEY,
        "X-RapidAPI-Host": "jsearch.p.rapidapi.com"
    }
    try:
        response = requests.get(url, headers=headers, params=querystring, timeout=10)
        if response.status_code == 200:
            data = response.json().get("data", [])
            jobs = []
            for j in data[:limit]:
                jobs.append({
                    "title": j.get("job_title", "Untitled"),
                    "company": j.get("employer_name", "Unknown"),
                    "location": j.get("job_city", "") + ", " + j.get("job_country", ""),
                    "url": j.get("job_apply_link", j.get("job_google_link", "#")),
                    "description": j.get("job_description", ""),
                    "date_display": f"📅 Posted: {j.get('job_posted_at_datetime_utc', 'recently')[:10]}",
                    "closing_date": None,
                    "created": j.get('job_posted_at_datetime_utc', ''),
                    "is_expired": False
                })
            return jobs, {"raw_count": len(data), "active_count": len(jobs)}
        else:
            return [], {"error": f"HTTP {response.status_code}"}
    except Exception as e:
        return [], {"error": str(e)}

@st.cache_data(ttl=3600)
def get_alternative_job_titles(cv_text, primary_role):
    """
    Ask Gemini to suggest 3-5 real alternative job titles that match the candidate's profile.
    Returns a list of strings (max 5).
    """
    prompt = f"""
    You are a career expert. Based on the CV below and the primary job title '{primary_role}',
    suggest 3 to 5 alternative job titles that are commonly used in the industry and would be relevant for this candidate.
    
    IMPORTANT:
    - Only return real, existing job titles. Do NOT make up fake titles.
    - Return ONLY a JSON array of strings. Example: ["Accounting Clerk", "Finance Assistant", "Bookkeeper"]
    - Do not include explanations.
    
    CV:
    {cv_text[:6000]}
    """
    try:
        response = model.generate_content(prompt)
        raw = clean_json_response(response.text)
        alternatives = json.loads(raw)
        if isinstance(alternatives, list) and len(alternatives) > 0:
            unique_titles = []
            for title in alternatives:
                if title.lower() != primary_role.lower() and title not in unique_titles:
                    unique_titles.append(title)
            return unique_titles[:5]
    except:
        pass
    return []

def get_job_matches(cv_text, analysis, manual_query, country_name, country_code, location_refine, limit=5, use_alternatives=True):
    target_roles = analysis.get("target_roles", [])
    
    if manual_query and len(manual_query.strip()) > 2:
        query = manual_query.strip()
        source = "manual override"
    elif target_roles and target_roles[0] != "N/A":
        query = target_roles[0]
        source = "CV analysis"
    else:
        query = generate_job_query(cv_text).strip()
        source = "CV text extraction"
    
    if not query or len(query) < 3:
        query = "Business Analyst"
        source = "fallback"
        st.warning(f"⚠️ Using fallback job title: '{query}'")
    
    original_query = query
    if any(word in query.lower() for word in ["head", "director", "chief", "vp", "vice president", "senior director"]):
        words = query.split()
        if words[0].lower() in ["head", "director", "chief", "vp"]:
            broader = " ".join(words[1:]).strip()
            if broader and len(broader) > 3:
                query = broader
                st.info(f"🔍 Broadened search from '{original_query}' to '{query}' for more results.")
    
    if "," in query:
        query = query.split(",")[0].strip()
    
    # --- Get alternative titles from Gemini ---
    search_queries = [query]
    if use_alternatives and not manual_query:
        with st.spinner("🧠 Generating related job titles to expand search..."):
            alternatives = get_alternative_job_titles(cv_text, query)
            if alternatives:
                search_queries.extend(alternatives)
                st.success(f"🔎 Also searching for: {', '.join(alternatives)}")
    
    st.success(f"🎯 Primary search: **{query}** (from {source}) in {country_name}")
    
    adzuna_supported = ["us", "gb", "ca", "au", "de", "fr", "in", "za"]
    all_jobs = []
    
    # Determine per-query limit
    per_query_limit = max(3, limit // len(search_queries) + 1)
    
    for q in search_queries:
        if country_code in adzuna_supported:
            jobs, log_info = get_jobs_from_adzuna(q, country_code, location_refine, per_query_limit)
            if jobs:
                all_jobs.extend(jobs)
        elif RAPIDAPI_KEY:
            fallback_jobs, _ = get_jobs_from_jsearch(q, country_code, location_refine, per_query_limit)
            if fallback_jobs:
                all_jobs.extend(fallback_jobs)
    
    unique_jobs = deduplicate_jobs(all_jobs)[:limit]
    
    if not unique_jobs:
        st.warning(f"No active jobs found for any of the searched titles.")
        if country_code in adzuna_supported:
            _, log_info = get_jobs_from_adzuna(query, country_code, location_refine, 5)
            with st.expander("🔧 Adzuna API Diagnostics", expanded=False):
                st.write(f"**Status Code:** {log_info.get('status')}")
                st.write(f"**Raw jobs returned:** {log_info.get('raw_count')}")
                st.write(f"**Active jobs after date filter:** {log_info.get('active_count')}")
                if log_info.get('error'):
                    st.error(f"**Error:** {log_info['error']}")
        # Fallback links
        encoded_query = query.replace(' ', '+')
        encoded_location = location_refine.replace(' ', '+') if location_refine else country_code.upper()
        st.info(f"💡 **Try these real job boards:**\n"
                f"- [Indeed](https://www.indeed.com/jobs?q={encoded_query}&l={encoded_location})\n"
                f"- [LinkedIn](https://www.linkedin.com/jobs/search?keywords={query.replace(' ', '%20')}&location={location_refine or country_name})\n"
                f"- [Glassdoor](https://www.glassdoor.com/Job/jobs.htm?sc.keyword={encoded_query}&locT=C&locId=&locKeyword={encoded_location})\n"
                f"- Or adjust your job title (e.g., '{query}' → 'AI Researcher', 'Machine Learning Lead')")
    else:
        st.success(f"✅ Found {len(unique_jobs)} relevant job(s)")
    
    return unique_jobs

@st.cache_data(ttl=3600)
def score_job_match(cv_text, job_title, job_description=""):
    prompt = f"""
    Score 0-100 match between CV and job '{job_title}'.
    Return JSON: {{"score": int, "reason": "string (max 10 words)"}}
    CV snippet:
    {cv_text[:2000]}
    Job description:
    {job_description[:500]}
    """
    response = model.generate_content(prompt)
    raw = clean_json_response(response.text)
    try:
        result = json.loads(raw)
        return result.get("score", 50), result.get("reason", "Based on role alignment")
    except:
        return 50, "General alignment"

@st.cache_data(ttl=3600)
def get_missing_keywords_preview(cv_text):
    prompt = f"""
    From this CV, identify 3 high-impact keywords missing that would most improve interview chances.
    Return ONLY comma-separated keywords.
    CV:
    {cv_text[:5000]}
    """
    response = model.generate_content(prompt)
    return response.text.strip()

def generate_improved_cv(cv_text, target_role):
    prompt = f"""
    Rewrite this CV for a {target_role} role.
    DO NOT invent experience. DO NOT change facts.
    Improve bullet points, add achievement language.
    Return the complete rewritten CV as plain text.
    Original: {cv_text[:10000]}
    """
    response = model.generate_content(prompt)
    return response.text

def create_docx_from_text(text, title="Document"):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph("Review before submission.")
    for line in text.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def safe_encode(text):
    return text.encode('latin-1', 'ignore').decode('latin-1')

def generate_pdf_report(analysis_full):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, safe_encode("Executive Career Intelligence Report"), ln=True)
    pdf.set_font("Arial", "I", 10)
    pdf.cell(0, 10, safe_encode(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"), ln=True)
    pdf.ln(5)
    pdf.set_font("Arial", size=12)
    for k, v in analysis_full.items():
        if isinstance(v, list):
            v = ", ".join(v)
        pdf.multi_cell(0, 8, safe_encode(f"{k.replace('_', ' ').title()}: {v}"))
    return pdf.output(dest='S').encode('latin-1')

def generate_ats_checklist(analysis_full):
    checklist = "✅ ATS OPTIMIZATION CHECKLIST\n\n"
    checklist += "Missing Keywords to Add:\n" + "\n".join(f"  • {kw}" for kw in analysis_full.get('missing_keywords', [])) + "\n\n"
    checklist += "Rewrite Suggestions:\n" + "\n".join(f"  • {sug}" for sug in analysis_full.get('rewrite_suggestions', [])) + "\n\n"
    return checklist

def remove_background_and_make_transparent(image_bytes):
    img = Image.open(BytesIO(image_bytes)).convert("RGBA")
    data = np.array(img)
    if data.shape[2] == 3:
        alpha = np.ones((data.shape[0], data.shape[1]), dtype=np.uint8) * 255
        white_mask = (data[:, :, 0] > 200) & (data[:, :, 1] > 200) & (data[:, :, 2] > 200)
        alpha[white_mask] = 0
        data = np.dstack((data, alpha))
    else:
        alpha = data[:, :, 3]
        white_mask = (data[:, :, 0] > 200) & (data[:, :, 1] > 200) & (data[:, :, 2] > 200)
        alpha[white_mask] = 0
        data[:, :, 3] = alpha
    result_img = Image.fromarray(data, "RGBA")
    output = BytesIO()
    result_img.save(output, format="PNG")
    output.seek(0)
    return output

@st.cache_data(ttl=3600)
def generate_job_description(job_title, company):
    prompt = f"Write a one‑sentence (max 150 characters) job description for a {job_title} position at {company}."
    try:
        return model.generate_content(prompt).text.strip()
    except:
        return "Description not available."

def generate_job_specific_cover_letter(cv_text, job_title, company, job_description):
    prompt = f"""
    Write a professional, tailored cover letter for the following job:

    Job Title: {job_title}
    Company: {company}
    Job Description: {job_description[:1500]}

    Use the candidate's CV below. DO NOT invent experience.
    Keep the letter to 250-350 words. Be specific about why the candidate is a good fit.

    CV:
    {cv_text[:6000]}
    """
    response = model.generate_content(prompt)
    return response.text

# ---------------------------
# 5. Intro Page
# ---------------------------
def intro_page():
    st.markdown("""
    <div class="hero">
        <div class="hero-badge">🤖 AI-POWERED CAREER INTELLIGENCE</div>
        <h1>📈 AI Career Intelligence</h1>
        <p>Get recruiter feedback, improve your CV, and apply to matched jobs – all in one platform.</p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("""
        <div style="background: white; border-radius: 20px; padding: 1.5rem; text-align: center; border: 1px solid #e2e8f0; height: 100%; color: #0f172a;">
            <div style="font-size: 1.5rem; font-weight: 700; color: #0f172a;">Free</div>
            <div style="font-size: 2rem; font-weight: 800; color: #4A90E2; margin: 1rem 0;">$0</div>
            <div style="text-align: left; margin: 1rem 0; color: #0f172a;">
                <div>✅ Basic CV scores</div>
                <div>✅ 1 job match</div>
                <div>✅ Preview of improvements</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Start Free", key="start_free", use_container_width=True):
            st.session_state.premium = False
            st.session_state.pro = False
            st.session_state.page = "workspace"
            st.rerun()

    with col2:
        st.markdown("""
        <div style="background: white; border-radius: 20px; padding: 1.5rem; text-align: center; border: 1px solid #e2e8f0; height: 100%; color: #0f172a;">
            <div style="display: inline-block; background-color: #6C63FF; color: white; padding: 0.2rem 1rem; border-radius: 30px; font-size: 0.7rem; margin-bottom: 1rem;">⭐ POPULAR</div>
            <div style="font-size: 1.5rem; font-weight: 700; color: #0f172a;">Premium</div>
            <div style="font-size: 2rem; font-weight: 800; color: #4A90E2; margin: 1rem 0;">$7<span style="font-size:1rem;">/month</span></div>
            <div style="font-size: 1.2rem; margin-bottom: 1rem; color: #0f172a;">or $29 lifetime</div>
            <div style="text-align: left; margin: 1rem 0; color: #0f172a;">
                <div>✅ Recruiter verdict</div>
                <div>✅ Missing keywords & rewrite suggestions</div>
                <div>✅ 10 job matches + match scores</div>
                <div>✅ Full cover‑letter diagnostics</div>
                <div>✅ ATS checklist & PDF report</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Choose Premium", key="choose_premium", use_container_width=True):
            st.session_state.premium = True
            st.session_state.pro = False
            st.session_state.page = "workspace"
            st.rerun()

    with col3:
        st.markdown("""
        <div style="background: white; border-radius: 20px; padding: 1.5rem; text-align: center; border: 1px solid #e2e8f0; height: 100%; color: #0f172a;">
            <div style="display: inline-block; background-color: #6C63FF; color: white; padding: 0.2rem 1rem; border-radius: 30px; font-size: 0.7rem; margin-bottom: 1rem;">🚀 BEST VALUE</div>
            <div style="font-size: 1.5rem; font-weight: 700; color: #0f172a;">Pro</div>
            <div style="font-size: 2rem; font-weight: 800; color: #4A90E2; margin: 1rem 0;">$15<span style="font-size:1rem;">/month</span></div>
            <div style="font-size: 1.2rem; margin-bottom: 1rem; color: #0f172a;">or $49 lifetime</div>
            <div style="text-align: left; margin: 1rem 0; color: #0f172a;">
                <div>✅ All Premium features</div>
                <div>✅ CV draft generator</div>
                <div>✅ Cover letter generator</div>
                <div>✅ Signature cleaner</div>
                <div>✅ 25+ job matches</div>
                <div>✅ Executive intelligence report</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        if st.button("Choose Pro", key="choose_pro", use_container_width=True):
            st.session_state.premium = False
            st.session_state.pro = True
            st.session_state.page = "workspace"
            st.rerun()

    st.markdown("---")
    with st.expander("🔓 Already have an unlock code?"):
        col_code1, col_code2 = st.columns(2)
        with col_code1:
            premium_input = st.text_input("Premium unlock code", type="password", key="intro_premium_code")
            if st.button("Apply Premium Code", key="intro_apply_premium"):
                if premium_input.strip() == PREMIUM_UNLOCK_CODE:
                    st.session_state.premium = True
                    st.session_state.pro = False
                    st.session_state.page = "workspace"
                    st.success("✅ Premium unlocked! Redirecting...")
                    st.rerun()
                else:
                    st.error("❌ Invalid Premium code.")
        with col_code2:
            pro_input = st.text_input("Pro unlock code", type="password", key="intro_pro_code")
            if st.button("Apply Pro Code", key="intro_apply_pro"):
                if pro_input.strip() == PRO_UNLOCK_CODE:
                    st.session_state.premium = False
                    st.session_state.pro = True
                    st.session_state.page = "workspace"
                    st.success("✅ Pro unlocked! Redirecting...")
                    st.rerun()
                else:
                    st.error("❌ Invalid Pro code.")

# ---------------------------
# 6. Workspace Page
# ---------------------------
def workspace_page():
    if st.session_state.pro:
        st.info("🚀 **Pro Tier Active** – Full application engine unlocked")
    elif st.session_state.premium:
        st.info("⭐ **Premium Tier Active** – Improvement tools unlocked")
    else:
        st.info("🔓 **Free Tier** – Basic scores and 1 job match. Upgrade to unlock more.")

    if not st.session_state.premium and not st.session_state.pro:
        with st.expander("🔓 Enter unlock code here", expanded=False):
            col_code1, col_code2 = st.columns(2)
            with col_code1:
                premium_input = st.text_input("Premium unlock code", type="password", key="top_premium_code")
                if st.button("Apply Premium Code", key="top_apply_premium"):
                    if premium_input.strip() == PREMIUM_UNLOCK_CODE:
                        st.session_state.premium = True
                        st.session_state.pro = False
                        st.success("✅ Premium unlocked! Refreshing...")
                        st.rerun()
                    else:
                        st.error("❌ Invalid Premium code.")
            with col_code2:
                pro_input = st.text_input("Pro unlock code", type="password", key="top_pro_code")
                if st.button("Apply Pro Code", key="top_apply_pro"):
                    if pro_input.strip() == PRO_UNLOCK_CODE:
                        st.session_state.premium = False
                        st.session_state.pro = True
                        st.success("✅ Pro unlocked! Refreshing...")
                        st.rerun()
                    else:
                        st.error("❌ Invalid Pro code.")

    uploaded_file = st.file_uploader("Upload your CV (PDF or DOCX)", type=["pdf", "docx"])
    if not uploaded_file:
        st.info("👆 Please upload your CV to begin.")
        st.stop()

    cv_text = extract_text_from_file(uploaded_file)
    st.session_state.cv_text = cv_text

    with st.spinner("Analysing your CV with AI..."):
        analysis = analyze_cv_cached(cv_text, full=False)
        st.session_state.analysis = analysis
        st.session_state.target_roles = analysis.get('target_roles', [])
        st.session_state.primary_role = st.session_state.target_roles[0] if st.session_state.target_roles and st.session_state.target_roles[0] != "N/A" else "Business Analyst"

    st.subheader("📝 Recruiter's Verdict on Your CV")
    st.info(f"**{analysis['recruiter_verdict']}**")

    strength = analysis['strength_score']
    if strength >= 70:
        interpretation = "✅ Your CV is competitive. Minor improvements could increase interview chances significantly."
    elif strength >= 50:
        interpretation = "📈 Your CV has good foundations. Addressing keyword gaps will boost recruiter interest."
    else:
        interpretation = "⚠️ Your CV needs structural improvement. The suggestions below will help you stand out."
    st.info(interpretation)

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{analysis['strength_score']}</div>
            <div class="metric-label">Overall CV Strength</div>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{analysis['ats_score']}</div>
            <div class="metric-label">ATS Readiness</div>
        </div>
        """, unsafe_allow_html=True)
    with col3:
        interview_label = analysis.get('interview_likelihood', 'Moderate')
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">{interview_label}</div>
            <div class="metric-label">Interview Likelihood</div>
        </div>
        """, unsafe_allow_html=True)
    with col4:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-value">✓</div>
            <div class="metric-label">Recruiter Verdict</div>
            <div style="font-size:0.75rem;">{analysis['recruiter_verdict'][:60]}...</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    col_left, col_mid, col_right = st.columns(3)

    # Improve CV card
    with col_left:
        with st.container():
            st.markdown('<div class="action-card">', unsafe_allow_html=True)
            st.markdown('<div class="action-title">📝 CV Mistakes & Improvements</div>', unsafe_allow_html=True)
            strengths = analysis.get('top_strengths', [])
            weaknesses = analysis.get('top_weaknesses', [])
            st.markdown("**Strengths:**")
            for s in strengths[:2]:
                st.markdown(f"- {s}")
            st.markdown("**Weaknesses (areas to fix):**")
            for w in weaknesses[:2]:
                st.markdown(f"- {w}")
            if not st.session_state.premium and not st.session_state.pro:
                st.markdown("🔒 **Upgrade to see full list of missing keywords and rewrite suggestions.**")
            else:
                full_analysis = analyze_cv_cached(st.session_state.cv_text, full=True)
                st.markdown("**Missing ATS Keywords:**")
                st.markdown(", ".join(full_analysis.get('missing_keywords', [])))
                st.markdown("**Rewrite Suggestions:**")
                for sug in full_analysis.get('rewrite_suggestions', []):
                    st.markdown(f"- {sug}")
                if st.session_state.pro:
                    if st.button("📄 Generate Improved CV Draft", use_container_width=True):
                        improved = generate_improved_cv(st.session_state.cv_text, st.session_state.primary_role)
                        st.session_state.generated_cv = improved
                    if st.session_state.generated_cv:
                        st.text_area("Improved CV Draft", st.session_state.generated_cv, height=200)
                        docx_file = create_docx_from_text(st.session_state.generated_cv, "Improved CV")
                        st.download_button("📥 Download CV", docx_file, file_name="improved_cv.docx")
                else:
                    st.info("🚀 Upgrade to Pro for CV draft generator")
            st.markdown('</div>', unsafe_allow_html=True)

    # Find Jobs card
    with col_mid:
        with st.container():
            st.markdown('<div class="action-card">', unsafe_allow_html=True)
            st.markdown('<div class="action-title">🌍 Find Matching Jobs</div>', unsafe_allow_html=True)
            st.caption(f"🎯 Searching for: **{st.session_state.primary_role}**")
            col_loc1, col_loc2 = st.columns(2)
            with col_loc1:
                country_display = st.selectbox("Country", list(COUNTRY_MAP.keys()), index=0, key="country_select")
                country_code = COUNTRY_MAP[country_display]
            with col_loc2:
                location_refine = st.text_input("City / Region", placeholder="e.g., London, Nairobi, Remote", key="location_input")
            manual_query = st.text_input("Override job title (optional)", placeholder=f"Leave empty to use {st.session_state.primary_role}", key="manual_query_input")
            
            use_expansion = st.checkbox("Expand search with related job titles (AI)", value=True)
            show_debug = st.checkbox("Show API diagnostics", value=False)
            
            search_clicked = st.button("🔍 Search for Jobs", use_container_width=True, type="primary")

            if st.session_state.pro:
                job_limit = 25
            elif st.session_state.premium:
                job_limit = 10
            else:
                job_limit = 1

            if search_clicked:
                with st.spinner("Searching for jobs..."):
                    try:
                        jobs = get_job_matches(
                            st.session_state.cv_text,
                            st.session_state.analysis,
                            manual_query,
                            country_display,
                            country_code,
                            location_refine,
                            limit=job_limit,
                            use_alternatives=use_expansion
                        )
                        st.session_state.jobs = jobs
                        st.session_state.match_scores = {}
                        if not jobs:
                            st.warning("No active jobs found. Try a different country or job title.")
                    except Exception as e:
                        st.error(f"Job search failed: {e}")
                        st.session_state.jobs = []

            if st.session_state.jobs:
                for idx, job in enumerate(st.session_state.jobs):
                    with st.container():
                        st.markdown(f"**{job['title']}**")
                        st.markdown(f"*{job['company']}*")
                        col_date, col_loc = st.columns(2)
                        with col_date:
                            st.caption(job.get('date_display', '📅 Date not specified'))
                        with col_loc:
                            st.caption(f"📍 {job.get('location', 'Not specified')}")
                        
                        # Show closing date warning if near deadline
                        if job.get('closing_date'):
                            close_date = parse_adzuna_date(job['closing_date'])
                            if close_date:
                                days_left = (close_date.date() - datetime.now().date()).days
                                if days_left <= 7:
                                    st.warning(f"⚠️ Closing in {days_left} day{'s' if days_left != 1 else ''}!")

                        raw_desc = job.get('description', '')
                        clean_desc = re.sub(r'<[^>]+>', '', raw_desc)
                        clean_desc = re.sub(r'\s+', ' ', clean_desc).strip()
                        with st.expander("📄 View job description"):
                            st.write(clean_desc)

                        score_key = f"score_{idx}"
                        if st.session_state.premium or st.session_state.pro:
                            if st.button(f"🎯 Show Match Score", key=f"match_btn_{idx}"):
                                score, reason = score_job_match(st.session_state.cv_text, job['title'], raw_desc)
                                st.session_state.match_scores[score_key] = (score, reason)
                            if score_key in st.session_state.match_scores:
                                score, reason = st.session_state.match_scores[score_key]
                                st.write(f"**Match Score:** {score}%")
                                st.caption(f"📝 {reason}")
                        else:
                            st.caption("🔒 Match score available after upgrade")

                        if st.button(f"✉️ Generate Cover Letter for this job", key=f"cover_btn_{idx}"):
                            with st.spinner("Generating tailored cover letter..."):
                                if st.session_state.premium or st.session_state.pro:
                                    job_desc = job.get('description', '')
                                    if not job_desc or len(job_desc) < 20:
                                        job_desc = f"A {job['title']} position at {job['company']}."
                                    letter = generate_job_specific_cover_letter(
                                        st.session_state.cv_text,
                                        job['title'],
                                        job['company'],
                                        job_desc
                                    )
                                    st.session_state.cover_letter_for_job = letter
                                else:
                                    st.session_state.cover_letter_for_job = "Upgrade to Premium to generate cover letters."
                        if st.session_state.cover_letter_for_job:
                            st.text_area("Generated Cover Letter", st.session_state.cover_letter_for_job, height=250)
                            docx_file = create_docx_from_text(st.session_state.cover_letter_for_job, "Cover Letter")
                            st.download_button("📥 Download Cover Letter", docx_file, file_name="cover_letter.docx")

                        if st.button(f"💾 Save this job", key=f"save_{idx}"):
                            if not any(saved.get('url') == job['url'] for saved in st.session_state.saved_jobs):
                                st.session_state.saved_jobs.append({
                                    "title": job['title'],
                                    "company": job['company'],
                                    "url": job['url'],
                                    "location": job.get('location', ''),
                                    "date_display": job.get('date_display', ''),
                                    "applied": False,
                                    "note": ""
                                })
                                st.success("Job saved!")
                        st.markdown(f"[Apply Now]({job['url']})")
                        st.markdown("---")
            st.markdown('</div>', unsafe_allow_html=True)

    # Saved Jobs card
    with col_right:
        with st.container():
            st.markdown('<div class="action-card">', unsafe_allow_html=True)
            st.markdown('<div class="action-title">💾 Your Saved Jobs</div>', unsafe_allow_html=True)
            if not st.session_state.saved_jobs:
                st.info("Jobs you save will appear here. Click 'Save this job' on any job listing.")
            else:
                for i, saved in enumerate(st.session_state.saved_jobs):
                    st.markdown(f"**{saved['title']}** at {saved['company']}")
                    st.caption(f"📍 {saved.get('location', '')} | {saved.get('date_display', '')}")
                    applied = st.checkbox("Applied", key=f"applied_{i}", value=saved.get('applied', False))
                    note = st.text_input("Note", key=f"note_{i}", value=saved.get('note', ''))
                    saved['applied'] = applied
                    saved['note'] = note
                    if st.button(f"Remove", key=f"remove_{i}"):
                        st.session_state.saved_jobs.pop(i)
                        st.rerun()
                    st.markdown("---")
            st.markdown('</div>', unsafe_allow_html=True)

    # Cover Letter Feedback Section
    st.markdown("---")
    st.subheader("📝 Cover Letter Feedback")
    st.caption("Upload or paste your cover letter to get detailed recruiter feedback.")
    cl_tab1, cl_tab2 = st.tabs(["📁 Upload Cover Letter", "📝 Paste Cover Letter"])
    cover_letter_text = ""
    with cl_tab1:
        uploaded_cl = st.file_uploader("Upload cover letter (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"], key="cl_upload")
        if uploaded_cl:
            try:
                if uploaded_cl.name.endswith(".pdf"):
                    reader = PyPDF2.PdfReader(uploaded_cl)
                    cover_letter_text = "".join(page.extract_text() for page in reader.pages)
                elif uploaded_cl.name.endswith(".docx"):
                    doc = docx.Document(uploaded_cl)
                    cover_letter_text = "\n".join(para.text for para in doc.paragraphs)
                else:
                    cover_letter_text = uploaded_cl.read().decode("utf-8")
                st.success("✅ Cover letter loaded")
            except Exception as e:
                st.error(f"Error: {e}")
    with cl_tab2:
        cover_letter_text = st.text_area("Paste your cover letter", height=150, key="cl_paste")
    analyze_cl_clicked = st.button("🔍 Analyze Cover Letter", use_container_width=True, type="primary")
    if analyze_cl_clicked and cover_letter_text and len(cover_letter_text.strip()) > 50:
        with st.spinner("Evaluating your cover letter..."):
            if st.session_state.premium or st.session_state.pro:
                cl_analysis = analyze_cover_letter_full(cover_letter_text, st.session_state.primary_role)
                st.metric("Application Readiness Score", f"{cl_analysis.get('overall_score', 50)}/100")
                st.info(f"**Recruiter Feedback:** {cl_analysis.get('verdict', 'Review needed')}")
                st.markdown("**Detailed Dimensions:**")
                col_cl1, col_cl2, col_cl3, col_cl4 = st.columns(4)
                with col_cl1:
                    st.metric("Role Alignment", f"{cl_analysis.get('alignment_score', 50)}/100")
                with col_cl2:
                    st.metric("Personalization", f"{cl_analysis.get('personalization_score', 50)}/100")
                with col_cl3:
                    st.metric("Impact", f"{cl_analysis.get('impact_score', 50)}/100")
                with col_cl4:
                    st.metric("Structure", f"{cl_analysis.get('structure_score', 50)}/100")
                missing = cl_analysis.get('missing_elements', [])
                if missing:
                    st.markdown("**Missing Elements (what to improve):** " + ", ".join(missing[:3]))
            else:
                basic = review_cover_letter_basic(cover_letter_text, st.session_state.primary_role)
                st.metric("Cover Letter Score", f"{basic.get('overall_score', 50)}/100")
                st.info(f"**Feedback:** {basic.get('verdict', 'Review complete')}")
                st.markdown(f"**Improvement Preview:** {basic.get('top_missing', 'Needs stronger alignment')}")
                st.caption("🔒 **Upgrade to Premium for full dimension scores and detailed improvements**")
    elif analyze_cl_clicked and cover_letter_text:
        st.warning("Please provide a cover letter with at least 50 characters.")
    elif analyze_cl_clicked:
        st.warning("Please upload or paste a cover letter first.")

    # Upgrade & Reports
    st.markdown('<div id="upgrade"></div>', unsafe_allow_html=True)
    if not st.session_state.premium and not st.session_state.pro:
        st.markdown("---")
        st.subheader("🚀 Upgrade Your Career Toolkit")
        col_card1, col_card2 = st.columns(2)
        with col_card1:
            st.markdown("""
            <div style="background: white; border-radius: 20px; padding: 1.5rem; text-align: center; border: 1px solid #e2e8f0; color: #0f172a;">
                <div style="display: inline-block; background-color: #6C63FF; color: white; padding: 0.2rem 1rem; border-radius: 30px; font-size: 0.7rem; margin-bottom: 1rem;">⭐ MOST POPULAR</div>
                <div style="font-size: 1.5rem; font-weight: 700; color: #0f172a;">Premium</div>
                <div style="font-size: 2rem; font-weight: 800; color: #4A90E2; margin: 1rem 0;">$7<span style="font-size:1rem;">/month</span></div>
                <div style="font-size: 1.2rem; margin-bottom: 1rem; color: #0f172a;">or $29 lifetime</div>
                <div style="text-align: left; margin: 1rem 0; color: #0f172a;">
                    <div>✅ Recruiter verdict</div>
                    <div>✅ Missing keywords & rewrite suggestions</div>
                    <div>✅ 10 job matches + match scores</div>
                    <div>✅ Full cover‑letter diagnostics</div>
                    <div>✅ ATS checklist & PDF report</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("⭐ Premium Monthly $7", use_container_width=True):
                try:
                    session = stripe.checkout.Session.create(
                        payment_method_types=["card"],
                        line_items=[{"price": STRIPE_PRICE_ID_PREMIUM_MONTHLY, "quantity": 1}],
                        mode="subscription",
                        success_url=APP_URL + "?success_premium_monthly=true",
                        cancel_url=APP_URL,
                    )
                    st.markdown(f"<a href='{session.url}' target='_blank'>Pay securely</a>", unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Payment error: {e}")
            if st.button("⭐ Premium Lifetime $29", use_container_width=True):
                try:
                    session = stripe.checkout.Session.create(
                        payment_method_types=["card"],
                        line_items=[{"price": STRIPE_PRICE_ID_PREMIUM_LIFETIME, "quantity": 1}],
                        mode="payment",
                        success_url=APP_URL + "?success_premium_lifetime=true",
                        cancel_url=APP_URL,
                    )
                    st.markdown(f"<a href='{session.url}' target='_blank'>Pay securely</a>", unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Payment error: {e}")

        with col_card2:
            st.markdown("""
            <div style="background: white; border-radius: 20px; padding: 1.5rem; text-align: center; border: 1px solid #e2e8f0; color: #0f172a;">
                <div style="display: inline-block; background-color: #6C63FF; color: white; padding: 0.2rem 1rem; border-radius: 30px; font-size: 0.7rem; margin-bottom: 1rem;">🚀 BEST VALUE</div>
                <div style="font-size: 1.5rem; font-weight: 700; color: #0f172a;">Pro</div>
                <div style="font-size: 2rem; font-weight: 800; color: #4A90E2; margin: 1rem 0;">$15<span style="font-size:1rem;">/month</span></div>
                <div style="font-size: 1.2rem; margin-bottom: 1rem; color: #0f172a;">or $49 lifetime</div>
                <div style="text-align: left; margin: 1rem 0; color: #0f172a;">
                    <div>✅ All Premium features</div>
                    <div>✅ CV draft generator</div>
                    <div>✅ Cover letter generator</div>
                    <div>✅ Signature cleaner</div>
                    <div>✅ 25+ job matches</div>
                    <div>✅ Executive intelligence report</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("🚀 Pro Monthly $15", use_container_width=True):
                try:
                    session = stripe.checkout.Session.create(
                        payment_method_types=["card"],
                        line_items=[{"price": STRIPE_PRICE_ID_PRO_MONTHLY, "quantity": 1}],
                        mode="subscription",
                        success_url=APP_URL + "?success_pro_monthly=true",
                        cancel_url=APP_URL,
                    )
                    st.markdown(f"<a href='{session.url}' target='_blank'>Pay securely</a>", unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Payment error: {e}")
            if st.button("🚀 Pro Lifetime $49", use_container_width=True):
                try:
                    session = stripe.checkout.Session.create(
                        payment_method_types=["card"],
                        line_items=[{"price": STRIPE_PRICE_ID_PRO_LIFETIME, "quantity": 1}],
                        mode="payment",
                        success_url=APP_URL + "?success_pro_lifetime=true",
                        cancel_url=APP_URL,
                    )
                    st.markdown(f"<a href='{session.url}' target='_blank'>Pay securely</a>", unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Payment error: {e}")

        st.markdown("---")
        st.subheader("🔓 Already have a code?")
        col_code1, col_code2 = st.columns(2)
        with col_code1:
            premium_input = st.text_input("Premium unlock code", type="password", key="workspace_premium_code")
            if st.button("Apply Premium Code", key="workspace_apply_premium"):
                if premium_input.strip() == PREMIUM_UNLOCK_CODE:
                    st.session_state.premium = True
                    st.session_state.pro = False
                    st.success("✅ Premium unlocked! Refreshing...")
                    st.rerun()
                else:
                    st.error("❌ Invalid Premium code.")
        with col_code2:
            pro_input = st.text_input("Pro unlock code", type="password", key="workspace_pro_code")
            if st.button("Apply Pro Code", key="workspace_apply_pro"):
                if pro_input.strip() == PRO_UNLOCK_CODE:
                    st.session_state.premium = False
                    st.session_state.pro = True
                    st.success("✅ Pro unlocked! Refreshing...")
                    st.rerun()
                else:
                    st.error("❌ Invalid Pro code.")

        st.subheader("📄 Reports")
        st.info("🔒 **PDF report and ATS checklist are available after upgrading to Premium or Pro.**")

    elif st.session_state.premium and not st.session_state.pro:
        st.markdown("---")
        st.subheader("📈 Upgrade to Pro")
        st.info("You are on Premium. Upgrade to Pro to unlock CV draft generator, cover letter generator, signature cleaner, 25+ job matches, and executive report.")
        col_up_btn1, col_up_btn2 = st.columns(2)
        with col_up_btn1:
            if st.button("🚀 Upgrade to Pro Monthly $15", use_container_width=True):
                try:
                    session = stripe.checkout.Session.create(
                        payment_method_types=["card"],
                        line_items=[{"price": STRIPE_PRICE_ID_PRO_MONTHLY, "quantity": 1}],
                        mode="subscription",
                        success_url=APP_URL + "?success_pro_monthly=true",
                        cancel_url=APP_URL,
                    )
                    st.markdown(f"<a href='{session.url}' target='_blank'>Pay securely</a>", unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Payment error: {e}")
        with col_up_btn2:
            if st.button("🚀 Upgrade to Pro Lifetime $49", use_container_width=True):
                try:
                    session = stripe.checkout.Session.create(
                        payment_method_types=["card"],
                        line_items=[{"price": STRIPE_PRICE_ID_PRO_LIFETIME, "quantity": 1}],
                        mode="payment",
                        success_url=APP_URL + "?success_pro_lifetime=true",
                        cancel_url=APP_URL,
                    )
                    st.markdown(f"<a href='{session.url}' target='_blank'>Pay securely</a>", unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Payment error: {e}")
        st.markdown("---")
        st.subheader("📄 Reports")
        with st.spinner("Generating full analysis for report..."):
            full_analysis = analyze_cv_cached(st.session_state.cv_text, full=True)
        pdf_data = generate_pdf_report(full_analysis)
        st.download_button("📥 Download Executive PDF Report", pdf_data, file_name="executive_report.pdf")
        checklist_text = generate_ats_checklist(full_analysis)
        st.download_button("📋 Download ATS Optimization Checklist", checklist_text, file_name="ats_checklist.txt")

    else:  # Pro user
        st.markdown("---")
        st.subheader("📄 Reports")
        with st.spinner("Generating full analysis for report..."):
            full_analysis = analyze_cv_cached(st.session_state.cv_text, full=True)
        pdf_data = generate_pdf_report(full_analysis)
        st.download_button("📥 Download Executive PDF Report", pdf_data, file_name="executive_report.pdf")
        checklist_text = generate_ats_checklist(full_analysis)
        st.download_button("📋 Download ATS Optimization Checklist", checklist_text, file_name="ats_checklist.txt")

    if st.session_state.pro:
        st.subheader("✍️ Signature Cleaner")
        uploaded_sig = st.file_uploader("Upload signature image (JPG, PNG, or JPEG)", type=["jpg", "jpeg", "png"], key="sig_upload")
        if uploaded_sig:
            with st.spinner("Cleaning signature..."):
                try:
                    transparent = remove_background_and_make_transparent(uploaded_sig.read())
                    st.success("✅ Signature cleaned!")
                    st.image(transparent, width=200)
                    st.download_button("📥 Download Transparent PNG", transparent, file_name="signature_clean.png", mime="image/png")
                except Exception as e:
                    st.error(f"Error processing image: {e}")

# ---------------------------
# 7. Main Router
# ---------------------------
if st.session_state.page == "intro":
    intro_page()
else:
    workspace_page()